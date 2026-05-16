#!/usr/bin/env python3
"""将求职背调 Markdown 报告转为排版良好的 Word (.docx) 格式"""
import sys, os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

input_path = sys.argv[1] if len(sys.argv) > 1 else None
if not input_path:
    print("用法: python3 scripts/to-docx.py <报告文件.md>")
    sys.exit(1)

resolved = os.path.expanduser(input_path)
if not os.path.exists(resolved):
    print(f"文件不存在: {resolved}")
    sys.exit(1)

with open(resolved, "r", encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Heading styles
for lvl, (size, color_hex) in {1: (22, '1a1a1a'), 2: (15, '333333'), 3: (12.5, '555555')}.items():
    h = doc.styles[f'Heading {lvl}']
    h.font.name = '微软雅黑'
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor(*tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4)))
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if lvl == 1:
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(12)
    elif lvl == 2:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
    else:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)

FONT_NAME = '微软雅黑'

def set_run_font(run, size=Pt(10.5), bold=False, color=None, italic=False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def parse_inline(text):
    """Parse **bold** and other inline marks, return list of (text, bold, italic) tuples"""
    result = []
    # Split by **pairs**
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            result.append((part[2:-2], True, False))
        elif part.startswith('`') and part.endswith('`'):
            result.append((part[1:-1], False, False))  # code style
        else:
            result.append((part, False, False))
    return result

def add_paragraph_with_bold(doc, text, font_size=Pt(10.5), indent_level=0):
    """Add a paragraph with inline bold support"""
    p = doc.add_paragraph()
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(indent_level * 0.75)
    parts = parse_inline(text)
    for txt, bold, _italic in parts:
        run = p.add_run(txt)
        set_run_font(run, size=font_size, bold=bold)
    return p

def add_cell_text(cell, text, bold=False, font_size=Pt(9)):
    """Set cell text with optional bold, clearing defaults"""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text.strip())
    set_run_font(run, size=font_size, bold=bold)

def add_table_from_md(doc, rows_data):
    """Create a proper Word table from markdown table rows"""
    if not rows_data or len(rows_data) < 1:
        return

    # Ensure all rows have same column count
    max_cols = max(len(r) for r in rows_data)
    for r in rows_data:
        while len(r) < max_cols:
            r.append('')

    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for j in range(max_cols):
        add_cell_text(table.rows[0].cells[j], rows_data[0][j], bold=True, font_size=Pt(9))

    # Data rows
    for i in range(1, len(rows_data)):
        for j in range(max_cols):
            add_cell_text(table.rows[i].cells[j], rows_data[i][j], bold=False, font_size=Pt(9))

    # Spacer
    doc.add_paragraph()

# --- Main parsing loop ---
i = 0
table_rows = []
in_code_block = False
in_blockquote = False
blockquote_lines = []

def flush_table():
    global table_rows
    if table_rows:
        # Filter out rows that are all separator chars (e.g., |-----|-----|)
        filtered = []
        for row in table_rows:
            if not all(re.match(r'^[\s\-:]+$', c) for c in row):
                filtered.append(row)
        if filtered:
            add_table_from_md(doc, filtered)
        table_rows = []

def flush_blockquote():
    global blockquote_lines
    if blockquote_lines:
        text = ' '.join(blockquote_lines)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text)
        set_run_font(run, size=Pt(9.5), color=(100, 100, 100), italic=True)
        blockquote_lines = []

while i < len(lines):
    line = lines[i].rstrip()

    # Empty line: flush accumulated state
    if not line:
        flush_table()
        flush_blockquote()
        i += 1
        continue

    # Code blocks
    if line.startswith('```'):
        in_code_block = not in_code_block
        i += 1
        continue
    if in_code_block:
        i += 1
        continue

    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        flush_table(); flush_blockquote()
        doc.add_heading(line[2:].strip(), level=1)
        i += 1
        continue
    if line.startswith('## ') and not line.startswith('### '):
        flush_table(); flush_blockquote()
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue
    if line.startswith('### '):
        flush_table(); flush_blockquote()
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue

    # Horizontal rule
    if line.strip() == '---':
        flush_table(); flush_blockquote()
        p = doc.add_paragraph()
        run = p.add_run('─' * 50)
        set_run_font(run, size=Pt(8), color=(180, 180, 180))
        i += 1
        continue

    # Blockquote (collect consecutive lines)
    if line.startswith('> '):
        flush_table()
        blockquote_lines.append(line[2:])
        # Check if next line is also blockquote
        if i + 1 < len(lines) and lines[i + 1].startswith('> '):
            i += 1
            continue
        flush_blockquote()
        i += 1
        continue

    # Table row
    if line.startswith('|') and line.endswith('|'):
        flush_blockquote()
        cells = [c.strip() for c in line.split('|')[1:-1]]
        table_rows.append(cells)
        i += 1
        continue

    # Non-table line after table rows → flush table
    if table_rows and not (line.startswith('|') and line.endswith('|')):
        flush_table()

    # Bullet list
    if re.match(r'^(\s*)[-*]\s', line):
        flush_table(); flush_blockquote()
        indent = len(re.match(r'^\s*', line).group())
        add_paragraph_with_bold(doc, re.sub(r'^\s*[-*]\s+', '', line),
                                indent_level=indent // 2)
        i += 1
        continue

    # Numbered list
    if re.match(r'^\s*\d+\.\s', line):
        flush_table(); flush_blockquote()
        add_paragraph_with_bold(doc, re.sub(r'^\s*\d+\.\s+', '', line), indent_level=1)
        i += 1
        continue

    # Regular paragraph
    line_clean = line.strip()
    if line_clean:
        add_paragraph_with_bold(doc, line_clean)

    i += 1

# Final flush
flush_table()
flush_blockquote()

out_path = resolved.replace('.md', '.docx')
doc.save(out_path)
print(f"Word 版已生成: {out_path}")
